# processing.py

from azure.storage.blob import BlobServiceClient
from azure.core.exceptions import HttpResponseError
from urllib.parse import quote
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document

# ★ 日本語フォントを正しく扱うために追加
from docx.oxml.ns import qn

from deepgram import Deepgram
import openai
import os
from dotenv import load_dotenv

# 必要な環境変数のロード
load_dotenv()

# storage.py の download_blob を使用（必要に応じて使用）
from storage import download_blob

# .env ファイルから環境変数を取得
AZURE_STORAGE_CONNECTION_STRING = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
AZURE_STORAGE_CONTAINER_NAME    = os.getenv("AZURE_STORAGE_CONTAINER_NAME")
OPENAI_API_KEY                  = os.getenv("OPENAI_API_KEY")
OPENAI_API_BASE                 = os.getenv("OPENAI_API_BASE")
DEPLOYMENT_ID                   = os.getenv("DEPLOYMENT_ID")
DEEPGRAM_API_KEY                = os.getenv("DEEPGRAM_API_KEY")

# OpenAI, Deepgram の設定
openai.api_key  = OPENAI_API_KEY
openai.api_base = OPENAI_API_BASE
openai.api_type = "azure"
openai.api_version = "2024-08-01-preview"

deepgram_client = Deepgram(DEEPGRAM_API_KEY)

# Azure Blob Storage クライアント (必要に応じて使用)
blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
container_client = blob_service_client.get_container_client(AZURE_STORAGE_CONTAINER_NAME)

# ローカル上の出力用フォルダ（固定名として定義している例）
UPLOAD_FOLDER = "./uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


async def process_files(audio_file_path, word_file_path, output_file_path):
    """
    音声ファイル (audio_file_path) と Word ファイル (word_file_path) を処理し、
    会議の情報を抽出して更新した Word ファイルを output_file_path に保存する。
    """
    try:
        # ============================
        # 1) Deepgram で音声解析
        # ============================
        with open(audio_file_path, "rb") as file:
            buffer_data = file.read()

        # Deepgram API オプション
        options = {
            "model": "nova-2-general",
            "detect_language": True,
            "diarize": True,
            "utterances": True,
        }

        # ファイル拡張子に応じて mimetype を推定
        if audio_file_path.endswith(".wav"):
            mimetype = "audio/wav"
        elif audio_file_path.endswith(".mp4"):
            mimetype = "audio/mp4"
        else:
            mimetype = "audio/mpeg"

        # 非同期の Deepgram API 呼び出し
        response = await deepgram_client.transcription.prerecorded(
            {"buffer": buffer_data, "mimetype": mimetype},
            options
        )
        print("Deepgram API のトランスクリプションが完了しました。")

        # 発話ごとの整形 (話者IDを含む)
        transcription = "\n".join(
            f"[Speaker {u['speaker']}] {u['transcript']}"
            for u in response["results"]["utterances"]
        )

        # ============================
        # 2) OpenAI で話者名を推測し、連続発話を結合
        # ============================
        prompt = (
            "以下のトランスクリプションを基に、話者を名前で推測し、"
            "同じ話者の連続する発言を結合してください。\n\n"
            f"トランスクリプション:\n{transcription}\n\n"
            "出力形式:\n"
            "[<話者名>] 発言内容\n"
            "[<話者名>] 発言内容\n"
        )

        response = openai.ChatCompletion.create(
            engine=DEPLOYMENT_ID,
            messages=[
                {
                    "role": "system",
                    "content": "あなたは話者を特定し、トランスクリプションを整理するアシスタントです。"
                },
                {
                    "role": "user",
                    "content": prompt
                },
            ],
            max_tokens=2000,
            temperature=0,
        )
        updated_transcription = response["choices"][0]["message"]["content"]
        print("OpenAI によるトランスクリプション整理が完了しました。")

        # ============================
        # 3) 議事録情報を抽出
        # ============================
        extracted_info = await extract_meeting_info_from_transcription(updated_transcription)

        # ============================
        # 4) Word ファイルを更新 & 保存
        # ============================
        # ここで output_file_path に保存する
        write_information_to_existing_table(
            word_file_path,
            output_file_path,
            extracted_info,
            updated_transcription
        )

    except Exception as e:
        # エラー内容をまとめて通知
        raise Exception(f"音声処理中にエラーが発生しました: {e}")


async def extract_meeting_info_from_transcription(transcription):
    """
    整理されたトランスクリプションから主要な会議情報を抽出
    (出席者, 次回予定, 次回開催場所, 決定事項, 宿題事項)
    """
    prompt = (
        "以下の議事録から、次の情報を抽出してください:\n\n"
        "1. 出席者\n"
        "2. 次回予定\n"
        "3. 次回開催場所\n"
        "4. 決定事項\n"
        "5. 宿題事項\n\n"
        f"議事録:\n{transcription}\n\n"
        "出力形式:\n"
        "出席者: <内容>\n"
        "次回予定: <内容>\n"
        "次回開催場所: <内容>\n"
        "決定事項:\n- <内容>\n"
        "宿題事項:\n- <内容>\n"
    )

    try:
        response = openai.ChatCompletion.create(
            engine=DEPLOYMENT_ID,
            messages=[
                {
                    "role": "system",
                    "content": "あなたは議事録を解析するアシスタントです。"
                },
                {
                    "role": "user",
                    "content": prompt
                },
            ],
            max_tokens=1000,
            temperature=0,
        )
        content = response["choices"][0]["message"]["content"]

        # 情報を格納する辞書
        extracted_info = {
            "出席者": "",
            "次回予定": "",
            "次回開催場所": "",
            "決定事項": "",
            "宿題事項": "",
        }

        current_section = None
        for line in content.splitlines():
            line = line.strip()
            if line.startswith("出席者:"):
                current_section = "出席者"
                extracted_info[current_section] = line.replace("出席者:", "").strip()
            elif line.startswith("次回予定:"):
                current_section = "次回予定"
                extracted_info[current_section] = line.replace("次回予定:", "").strip()
            elif line.startswith("次回開催場所:"):
                current_section = "次回開催場所"
                extracted_info[current_section] = line.replace("次回開催場所:", "").strip()
            elif line.startswith("決定事項:"):
                current_section = "決定事項"
            elif line.startswith("宿題事項:"):
                current_section = "宿題事項"
            elif line.startswith("-") and current_section:
                extracted_info[current_section] += line + "\n"

        # 空欄なら「なし」に置き換え
        for key in extracted_info:
            extracted_info[key] = extracted_info[key].strip() or "なし"

        return extracted_info

    except Exception as e:
        raise Exception(f"情報抽出中にエラーが発生しました: {e}")


def write_information_to_existing_table(word_file, output_path, extracted_info, transcription):
    """
    Word ファイルを更新し、情報をテーブルに書き込み。
    さらに、文書全体のフォントを 'Meiryo' に統一して保存する。
    """
    try:
        doc = Document(word_file)

        # -------------------------
        # 1) テーブルの各セルを更新
        # -------------------------
        for table in doc.tables:
            for row in table.rows:
                first_cell_text = row.cells[0].text
                if "出席者" in first_cell_text:
                    row.cells[1].text = extracted_info.get("出席者", "なし")
                elif "次回予定" in first_cell_text:
                    row.cells[1].text = extracted_info.get("次回予定", "なし")
                elif "次回開催場所" in first_cell_text:
                    row.cells[1].text = extracted_info.get("次回開催場所", "なし")
                elif "決定事項" in first_cell_text:
                    row.cells[1].text = extracted_info.get("決定事項", "なし")
                elif "宿題事項" in first_cell_text:
                    row.cells[1].text = extracted_info.get("宿題事項", "なし")

        # -------------------------
        # 2) "■ 議事" 段落を探して追記
        # -------------------------
        for i, paragraph in enumerate(doc.paragraphs):
            if "■ 議事" in paragraph.text:
                # 不要行を削除
                cleaned_transcription = remove_line_starting_with_and_ending_with(
                    transcription,
                    start_substring="以下",
                    end_substring=":"
                )
                # 議事の段落に追記
                paragraph.text += "\n" + cleaned_transcription

                # 次の段落があれば "以上" を右寄せで追加
                if i + 1 < len(doc.paragraphs):
                    new_paragraph = doc.paragraphs[i + 1]
                    new_paragraph.text = "以上"
                    new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                break

        # -------------------------
        # 3) 文書全体のフォントを 'Meiryo' に統一
        # -------------------------
        def set_meiryo_font(run):
            """
            日本語が適切に表示されるように、東アジアフォントも Meiryo に設定。
            """
            run.font.name = 'Meiryo'
            # ★ ここがポイント：東アジア文字も同じフォントに
            run.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')

        # (A) 段落 (doc.paragraphs)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                set_meiryo_font(run)

        # (B) テーブル (cells 内の paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_meiryo_font(run)

        # -------------------------
        # 4) 保存
        # -------------------------
        doc.save(output_path)

    except Exception as e:
        raise Exception(f"Word ファイルの更新中にエラーが発生しました: {e}")


def remove_line_starting_with_and_ending_with(text, start_substring, end_substring):
    """
    指定した文字列で始まり、かつ指定した文字列で終わる行を削除して返す。
    例: start_substring="以下", end_substring=":" の場合、
        "以下 ... :" で始まり終わる行を除外する。
    """
    lines = text.split("\n")
    filtered_lines = []
    for line in lines:
        # 先頭と末尾が指定文字列ならスキップ
        if line.startswith(start_substring) and line.endswith(end_substring):
            continue
        filtered_lines.append(line)
    return "\n".join(filtered_lines)
