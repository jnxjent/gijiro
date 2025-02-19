# docwriter.py

import re
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def write_information_to_existing_table(
    word_file_path: str,
    output_file_path: str,
    extracted_info: dict,
    full_transcription: str
):
    """
    1) テンプレートWordファイルを開く
    2) テーブル各セルへ抽出情報を書き込み。
       - '出席者' セルには "推定話者" の名前一覧を「、」区切りで1行にまとめる
    3) 「■議事」段落に全文文字起こしを追記 → "Speaker0→○○" の置換
    4) フォントをMeiryoに統一して保存
    """

    doc = Document(word_file_path)

    # 推定話者情報 (例: "- Speaker0 -> 本田\n- Speaker1 -> 橋本")
    speaker_info_text = extracted_info.get("推定話者", "")
    # => 辞書 {"Speaker0":"本田","Speaker1":"橋本"} を作成
    speaker_map = create_speaker_map(speaker_info_text)

    # ★ 出席者欄に書くため、「名前リスト」をパースする (敬称は外す)
    speaker_name_list = parse_speaker_names_only(speaker_info_text)
    # 例: ["本田", "橋本"]

    # ------------------------------------------------------
    # STEP A: テーブルの各セルを更新
    # ------------------------------------------------------
    for table in doc.tables:
        for row in table.rows:
            first_cell_text = row.cells[0].text.strip()

            if "出席者" in first_cell_text:
                # 出席者は「推定話者一覧」を「、」区切りで1行にまとめる
                if speaker_name_list:
                    row.cells[1].text = "、".join(speaker_name_list)
                else:
                    row.cells[1].text = "なし"

            elif "次回予定" in first_cell_text:
                row.cells[1].text = extracted_info.get("次回予定", "なし")

            elif "次回開催場所" in first_cell_text:
                row.cells[1].text = extracted_info.get("次回開催場所", "なし")

            elif "決定事項" in first_cell_text:
                row.cells[1].text = extracted_info.get("決定事項", "なし")

            elif "宿題事項" in first_cell_text:
                row.cells[1].text = extracted_info.get("宿題事項", "なし")

            elif "推定話者" in first_cell_text:
                # 別途「推定話者」という欄があるなら原文を記載
                row.cells[1].text = speaker_info_text or "不明"

    # ------------------------------------------------------
    # STEP B: 「■議事」段落を探して、全文文字起こしを追記
    #         → Speaker置換
    # ------------------------------------------------------
    paragraph_index = None
    for i, paragraph in enumerate(doc.paragraphs):
        if "■ 議事" in paragraph.text:
            paragraph_index = i
            # (1) そのまま追記（AIに再度かけない）
            paragraph.text += "\n" + full_transcription

            # (2) 次の段落に「以上」を右寄せで追加
            if i + 1 < len(doc.paragraphs):
                new_paragraph = doc.paragraphs[i + 1]
                new_paragraph.text = "以上"
                new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            break

    # 追記後、置換を実施
    if paragraph_index is not None:
        paragraph = doc.paragraphs[paragraph_index]
        replaced_paragraph_text = apply_speaker_map(paragraph.text, speaker_map)
        paragraph.text = replaced_paragraph_text

    # ------------------------------------------------------
    # STEP C: 文書全体のフォントを 'Meiryo' に統一
    # ------------------------------------------------------
    def set_meiryo_font(run):
        run.font.name = 'Meiryo'
        run.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Meiryo')

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_meiryo_font(run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_meiryo_font(run)

    # ------------------------------------------------------
    # STEP D: 保存
    # ------------------------------------------------------
    doc.save(output_file_path)


def create_speaker_map(speaker_info_text: str) -> dict:
    """
    例:
      - Speaker0 -> 本田
      - Speaker1 -> 橋本 さん
    のような文字列から、{"Speaker0": "本田", "Speaker1": "橋本"} を返す。
    
    改善点:
      - 行頭のハイフンや空白を柔軟に許容
      - "->" の前後の空白をトリム
      - 名前の末尾に付く「さん」や余計な空白を除去
    """
    mapping = {}
    # speaker_key と名前部分を抽出。名前部分は「さん」など余計な語尾を除去できるようにする
    pattern = re.compile(
        r"^-?\s*(?P<speaker>Speaker\s*\d+)\s*->\s*(?P<name>.+)$", re.IGNORECASE)
    
    for line in speaker_info_text.splitlines():
        line = line.strip()
        match = pattern.match(line)
        if match:
            speaker_key = match.group("speaker").strip()
            name_val = match.group("name").strip()
            # 名前の末尾の「さん」や不要な記号・空白を除去（必要に応じて他の敬称も追加可能）
            name_val = re.sub(r"[さん\s]+$", "", name_val)
            mapping[speaker_key] = name_val
    return mapping


def parse_speaker_names_only(speaker_info_text: str) -> list:
    """
    例:
      - Speaker0 -> 本田
      - Speaker1 -> 橋本 さん
    から ["本田", "橋本"] のリストを返す。
    """
    names = []
    # 同じ正規表現で抽出
    pattern = re.compile(
        r"^-?\s*Speaker\s*\d+\s*->\s*(?P<name>.+)$", re.IGNORECASE)
    for line in speaker_info_text.splitlines():
        line = line.strip()
        match = pattern.match(line)
        if match:
            name_val = match.group("name").strip()
            name_val = re.sub(r"[さん\s]+$", "", name_val)
            names.append(name_val)
    return names


def apply_speaker_map(text: str, speaker_map: dict) -> str:
    """
    本文中の "Speaker0" や "[Speaker 0]" などを、対応する名前に置換する。
    例:
      "Hello Speaker0 and [Speaker 0]!"  -> "Hello [本田] and [本田]!"
    
    改善点:
      - 正規表現で余分な空白や角括弧の有無を許容
    """
    replaced_text = text
    for speaker_key, name_val in speaker_map.items():
        # speaker_key から数字のみを抽出（例: "Speaker 0" -> "0"）
        digits = re.sub(r"\D", "", speaker_key)
        # 例: パターンは "[?]speaker\s*digits[?]" を許容する
        pattern = re.compile(rf"\[?\s*speaker\s*{digits}\s*\]?", re.IGNORECASE)
        replaced_text = pattern.sub(f"[{name_val}]", replaced_text)
    return replaced_text
